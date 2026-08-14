#!/usr/bin/env python3
"""Build PIM Architecture for Business Central as a Word document."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent
DOC_PATH = OUT_DIR / "PIM-Architecture.docx"
IMG_DIR = OUT_DIR / "_arch_images"

TEAL = "#004c3f"
GREEN = "#008060"
BG = "#f4f7f6"
LINE = "#c5d0cc"
TEXT = "#1a1a1a"
MUTED = "#5c6664"
WHITE = "#ffffff"
RGB_TEAL = RGBColor(0x00, 0x4C, 0x3F)
RGB_GREEN = RGBColor(0x00, 0x80, 0x60)
RGB_MUTED = RGBColor(0x5C, 0x66, 0x64)
RGB_BLACK = RGBColor(0x1A, 0x1A, 0x1A)


def rounded(ax, x, y, w, h, text, fc=WHITE, ec=TEAL, fontsize=8.5, fw="medium", tc=TEXT):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.2,
        facecolor=fc,
        edgecolor=ec,
        mutation_aspect=0.4,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        color=tc,
        fontweight=fw,
        wrap=True,
        family="DejaVu Sans",
    )
    return box


def arrow(ax, x1, y1, x2, y2, color=GREEN):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.3,
            color=color,
            shrinkA=2,
            shrinkB=2,
        )
    )


def new_fig(w=12.2, h=6.4):
    fig, ax = plt.subplots(figsize=(w, h), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.2)
    ax.axis("off")
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    return fig, ax


def save_fig(fig, name: str) -> Path:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / name
    fig.tight_layout(pad=0.3)
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def diagram_context() -> Path:
    fig, ax = new_fig(12.4, 6.8)
    ax.set_ylim(0, 6.6)
    ax.add_patch(
        FancyBboxPatch((0.25, 2.55), 5.7, 3.7, boxstyle="round,pad=0.04,rounding_size=0.1", facecolor=BG, edgecolor=TEAL, linewidth=1.4)
    )
    ax.text(3.1, 6.0, "MASTER PRODUCT COMPANY", ha="center", fontsize=10, color=TEAL, fontweight="bold", family="DejaVu Sans")
    rounded(ax, 0.55, 4.55, 5.1, 0.85, "Item  |  same Item No. across countries", fc=WHITE, fontsize=9)
    rounded(ax, 0.55, 3.5, 5.1, 0.85, "PIM model  |  family · attributes · values · category", fc=WHITE, fontsize=9)
    rounded(ax, 0.55, 2.75, 2.4, 0.6, "Webshop preview", fc=GREEN, ec=GREEN, tc=WHITE, fontsize=8.5, fw="bold")
    rounded(ax, 3.15, 2.75, 2.5, 0.6, "Company Sync", fc=TEAL, ec=TEAL, tc=WHITE, fontsize=8.5, fw="bold")

    ax.add_patch(
        FancyBboxPatch((6.25, 2.55), 5.5, 3.7, boxstyle="round,pad=0.04,rounding_size=0.1", facecolor="#eef6f3", edgecolor=GREEN, linewidth=1.4)
    )
    ax.text(9.0, 6.0, "CHILD COMPANIES  (same environment)", ha="center", fontsize=10, color=TEAL, fontweight="bold", family="DejaVu Sans")
    codes = [("DE", "Germany"), ("AT", "Austria"), ("ES", "Spain"), ("CH", "Switzerland"), ("NP", "Nonpa"), ("CZ", "Czech Republic")]
    for i, (code, name) in enumerate(codes):
        col, row = i % 3, i // 3
        rounded(ax, 6.5 + col * 1.7, 4.55 - row * 1.15, 1.5, 0.9, f"{code}\n{name}", fontsize=8)

    arrow(ax, 5.65, 3.05, 6.35, 3.05)
    ax.text(6.0, 3.22, "ChangeCompany", ha="center", fontsize=7, color=GREEN, family="DejaVu Sans")

    rounded(ax, 0.25, 0.35, 3.6, 1.7, "PIM / product user\nEnrich & publish", fontsize=9)
    rounded(ax, 4.2, 0.35, 3.6, 1.7, "Country sales / purchase\nUses same Item No.", fontsize=9)
    rounded(ax, 8.15, 0.35, 3.6, 1.7, "Internal storefront\nPIM data, not full ERP card", fontsize=9)
    ax.text(6.0, 0.08, "Separate tenants are out of scope — ChangeCompany works only inside one Business Central environment.", ha="center", fontsize=7.5, color=MUTED, family="DejaVu Sans")
    return save_fig(fig, "01-context.png")


def diagram_layers() -> Path:
    fig, ax = new_fig(12.2, 4.6)
    ax.set_ylim(0, 4.4)
    layers = [
        (0.3, "1  Setup", "Families\nAttributes\nCategories\nMarketplaces", GREEN),
        (3.25, "2  Enrichment", "Item + PIM fields\nProduct values\nCompleteness %\nPublish flag", TEAL),
        (6.2, "3  Distribution", "Company sync\nSame Item No.\nTemplate posting\nSync log", GREEN),
        (9.15, "4  Experience", "JSON payload\nControl add-in\nShopify-style UI\nPublished catalog", TEAL),
    ]
    for x, title, body, color in layers:
        rounded(ax, x, 2.35, 2.7, 0.7, title, fc=color, ec=color, tc=WHITE, fontsize=11, fw="bold")
        rounded(ax, x, 0.45, 2.7, 1.75, body, fontsize=9)
    for x in (3.0, 5.95, 8.9):
        arrow(ax, x, 2.7, x + 0.25, 2.7)
    ax.text(6.05, 3.85, "PIM does not replace ERP Item, inventory, costing, or documents", ha="center", fontsize=10, color=TEAL, fontweight="bold", family="DejaVu Sans")
    return save_fig(fig, "02-layers.png")


def diagram_model() -> Path:
    fig, ax = new_fig(12.4, 7.2)
    ax.set_ylim(0, 7.0)
    ax.set_xlim(0, 12.2)
    rounded(ax, 0.3, 5.7, 2.5, 0.9, "Attribute Group", fontsize=9, fw="bold")
    rounded(ax, 3.3, 5.7, 2.5, 0.9, "Attribute\n(+ Shopify Field)", fontsize=9, fw="bold")
    rounded(ax, 6.3, 5.7, 2.5, 0.9, "Attribute Option", fontsize=9, fw="bold")
    rounded(ax, 9.3, 5.7, 2.5, 0.9, "Category\n(parent tree)", fontsize=9, fw="bold")
    rounded(ax, 1.8, 3.85, 2.5, 0.9, "Family", fontsize=9, fw="bold", fc=BG)
    rounded(ax, 5.0, 3.85, 2.8, 0.9, "Family Attribute\n(required flag)", fontsize=9, fw="bold", fc=BG)
    rounded(ax, 3.3, 2.0, 3.4, 1.1, "ITEM\nFamily · Category · Published", fc=TEAL, ec=TEAL, tc=WHITE, fontsize=10, fw="bold")
    rounded(ax, 7.2, 2.05, 2.7, 1.0, "Product Value\nItem + Attribute", fontsize=9, fw="bold")
    rounded(ax, 0.4, 0.25, 3.5, 1.15, "Marketplace\nBC Company Name\nTemplate Item", fontsize=9, fw="bold", fc="#eef6f3", ec=GREEN)
    rounded(ax, 4.3, 0.25, 3.5, 1.15, "Item Marketplace\nSync ticks + status", fontsize=9, fw="bold", fc="#eef6f3", ec=GREEN)
    rounded(ax, 8.2, 0.25, 3.5, 1.15, "Sync Log", fontsize=9, fw="bold", fc="#eef6f3", ec=GREEN)
    arrow(ax, 2.8, 6.15, 3.3, 6.15)
    arrow(ax, 5.8, 6.15, 6.3, 6.15)
    arrow(ax, 3.05, 4.3, 3.3, 5.7)
    arrow(ax, 6.4, 4.75, 5.55, 5.7)
    arrow(ax, 3.05, 4.3, 4.5, 3.1)
    arrow(ax, 5.0, 2.55, 7.2, 2.55)
    arrow(ax, 5.0, 2.0, 6.05, 1.4)
    arrow(ax, 2.15, 1.4, 4.3, 0.85)
    return save_fig(fig, "03-data-model.png")


def diagram_sync() -> Path:
    fig, ax = new_fig(12.2, 6.4)
    steps = [
        (0.35, 4.55, "1. Sync action\non Item Card"),
        (3.25, 4.55, "2. Enabled\nmarketplaces"),
        (6.15, 4.55, "3. ChangeCompany\nto child"),
        (9.05, 4.55, "4. Same Item No.\nget or insert"),
        (0.35, 2.35, "5. Template in TARGET\nposting groups if new"),
        (3.25, 2.35, "6. Copy PIM +\nproduct master"),
        (6.15, 2.35, "7. Never copy\nstock / cost / docs"),
        (9.05, 2.35, "8. Write\nPIM Sync Log"),
    ]
    for x, y, t in steps:
        rounded(ax, x, y, 2.55, 1.25, t, fontsize=9)
    for x in (2.9, 5.8, 8.7):
        arrow(ax, x, 5.15, x + 0.35, 5.15)
    arrow(ax, 10.3, 4.55, 10.3, 3.7)
    arrow(ax, 10.3, 3.6, 1.6, 3.6)
    arrow(ax, 1.6, 3.6, 1.6, 3.6)
    arrow(ax, 1.62, 3.55, 1.62, 3.6)
    # down from 4 to 5 visually via left
    arrow(ax, 2.9, 2.95, 3.25, 2.95)
    arrow(ax, 5.8, 2.95, 6.15, 2.95)
    arrow(ax, 8.7, 2.95, 9.05, 2.95)
    ax.text(
        6.1,
        1.35,
        "Copied: description, GTIN, picture, variants, UOMs, translations,\nextended texts, documents, PIM values, item attributes by Name",
        ha="center",
        fontsize=8.5,
        color=TEAL,
        family="DejaVu Sans",
    )
    ax.text(6.1, 0.45, "Not copied: inventory, unit cost, sales/purchase orders, warehouse SKUs, TransferFields of whole Item", ha="center", fontsize=8.5, color="#a33", family="DejaVu Sans")
    return save_fig(fig, "04-sync.png")


def diagram_webshop() -> Path:
    fig, ax = new_fig(12.2, 5.2)
    ax.set_ylim(0, 5.0)
    boxes = [
        (0.3, 3.35, "User opens\nProduct Webshop"),
        (3.2, 3.35, "Control add-in\nControlReady"),
        (6.1, 3.35, "Product Visual Data\nJSON"),
        (9.0, 3.35, "SetProductData\nHTML storefront"),
        (0.3, 1.15, "Catalog:\nPIM Published only"),
        (3.2, 1.15, "Product page:\nPIM + variants + media"),
        (6.1, 1.15, "Shopify Field map\n(preview only)"),
        (9.0, 1.15, "Open Item Card\nfrom shop"),
    ]
    for x, y, t in boxes:
        rounded(ax, x, y, 2.65, 1.25, t, fontsize=9)
    for x in (2.95, 5.85, 8.75):
        arrow(ax, x, 3.95, x + 0.25, 3.95)
    arrow(ax, 10.3, 3.35, 10.3, 2.5)
    arrow(ax, 10.3, 2.45, 1.6, 2.45)
    arrow(ax, 1.6, 2.45, 1.6, 2.4)
    for x in (2.95, 5.85, 8.75):
        arrow(ax, x, 1.75, x + 0.25, 1.75)
    ax.text(6.1, 0.4, "Internal Business Central page — not a public website and not a live Shopify connection", ha="center", fontsize=8.5, color=MUTED, family="DejaVu Sans")
    return save_fig(fig, "05-webshop.png")


def set_run_font(run, name="Calibri", size=11, bold=False, color=RGB_BLACK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.replace("#", ""))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C5D0CC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_table(doc: Document, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "004C3F")
        set_cell_border(cell)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            if r % 2 == 1:
                shade_cell(cell, "F4F7F6")
            set_cell_border(cell)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGB_TEAL
        run.font.name = "Calibri"
    return p


def body(doc, text, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_run_font(run, size=9, color=RGB_MUTED)
    run.italic = True
    return p


def add_picture(doc, path: Path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.4))
    p.paragraph_format.space_after = Pt(4)


def set_cell_text(cell, text, size=11, bold=False, color=RGB_BLACK, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)


def build():
    imgs = {
        "context": diagram_context(),
        "layers": diagram_layers(),
        "model": diagram_model(),
        "sync": diagram_sync(),
        "webshop": diagram_webshop(),
    }

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("ICS Master  ·  PIM Architecture  ·  Confidential")
    set_run_font(run, size=8, color=RGB_MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("ZVG  ·  Dynamics 365 Business Central  ·  Page ")
    set_run_font(r, size=8, color=RGB_MUTED)
    add_page_number(fp)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ICS MASTER")
    set_run_font(r, size=14, bold=True, color=RGB_GREEN)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Best-practice PIM architecture")
    set_run_font(r, size=28, bold=True, color=RGB_TEAL)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Product Information Management in Microsoft Dynamics 365 Business Central")
    set_run_font(r, size=14, color=RGB_MUTED)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Akeneo-style enrichment  ·  Same Item No. across companies  ·  Shopify-style preview")
    set_run_font(r, size=11, color=RGB_GREEN)

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    facts = [
        ("Application", "ICS Master (publisher ZVG)"),
        ("Platform", "Dynamics 365 Business Central 27  ·  AL runtime 16"),
        ("Object ranges", "50350–50399 and 50600–50700"),
        ("Companies", "Master + DE, AT, ES, CH, NP, CZ"),
        ("Document type", "Solution architecture"),
    ]
    for i, (k, v) in enumerate(facts):
        set_cell_text(meta.rows[i].cells[0], k, bold=True, color=RGB_TEAL, size=11)
        set_cell_text(meta.rows[i].cells[1], v, size=11)
        shade_cell(meta.rows[i].cells[0], "F4F7F6")
        meta.rows[i].cells[0].width = Cm(4.5)
        meta.rows[i].cells[1].width = Cm(11.5)

    doc.add_page_break()

    heading(doc, "1. Executive summary", 1)
    body(
        doc,
        "This architecture puts Product Information Management (PIM) inside Business Central. "
        "It follows the Akeneo pattern (families, attributes, options, categories, completeness) without replacing the ERP Item. "
        "Sales and purchase in every country keep the same Item No. PIM stores commercial truth; the Item Card keeps posting, inventory, and documents.",
    )
    body(
        doc,
        "A master product company enriches items. PIM Company Sync copies product-master data into child companies in the same environment using ChangeCompany. "
        "An internal Shopify-style control add-in previews published products. This is not a public website and not a live Shopify connector.",
    )

    add_table(
        doc,
        ["Principle", "Decision"],
        [
            ["Identity", "Item No. is the same in master and in every country company."],
            ["Separation of concerns", "PIM is additive. Inventory, unit cost, and documents stay on standard Item."],
            ["Master → child", "Commercial / product-master data flows one way in this version."],
            ["Storefront", "Catalog shows only items with Published to Webshop."],
            ["PTE coexistence", "PIM does not use or overwrite PTE Marketplace Code / Master SKU."],
            ["Tenancy", "Sync works only when countries are companies in the same BC environment."],
        ],
        [4.2, 12.8],
    )

    heading(doc, "2. Context diagram", 1)
    body(doc, "People enrich in the master company. Child companies receive the same Item No. with commercial data. Country users continue to sell and buy on that number.")
    add_picture(doc, imgs["context"])
    caption(doc, "Figure 1 — Master product company, six child companies, and user roles")

    heading(doc, "3. Architecture layers", 1)
    add_picture(doc, imgs["layers"])
    caption(doc, "Figure 2 — Setup, enrichment, distribution, experience")

    add_table(
        doc,
        ["Layer", "Responsibility", "Main objects"],
        [
            ["1. Setup", "Reference data: families, attributes, groups, options, categories, marketplaces.", "Tables 50600–50605, 50636; pages 50610–50617, 50640; install 50634"],
            ["2. Enrichment", "Per-item family, values, completeness, publish flag, marketplace ticks.", "Item ext 50607; table 50606; pages 50618–50619; codeunit 50632; pageext 50353–50354"],
            ["3. Distribution", "Copy product master into child companies; log success and errors.", "Codeunit 50639; tables 50637–50638; pages 50641–50642"],
            ["4. Experience", "JSON storefront for published items (internal preview).", "Pages 50351–50352; codeunit 50633; control add-in ProductVisualViewer"],
        ],
        [3.0, 7.2, 6.8],
    )

    heading(doc, "4. Data model", 1)
    body(doc, "The model is Akeneo-shaped and sits next to Item. Only three fields are added on Item: PIM Family Code, PIM Category Code, and PIM Published. All attribute values live in PIM Product Value.")
    add_picture(doc, imgs["model"])
    caption(doc, "Figure 3 — PIM entities and their relationship to Item")

    heading(doc, "4.1 Akeneo mapping", 2)
    add_table(
        doc,
        ["Akeneo concept", "Business Central object", "Notes"],
        [
            ["Family", "PIM Family + PIM Family Attribute", "Required flags drive completeness %"],
            ["Attribute group", "PIM Attribute Group", "IDENT, MARKETING, SPECS, SEO in the default seed"],
            ["Attribute", "PIM Attribute", "Text, Number, Yes/No, Option, Date"],
            ["Attribute options", "PIM Attribute Option", "For example Color: Black, White, Grey, Blue, Red"],
            ["Category tree", "PIM Category", "Parent Code for hierarchy"],
            ["Product values", "PIM Product Value", "Keyed by Item No. + Attribute Code"],
            ["Completeness", "PIM Enrichment codeunit", "Filled required attributes / required × 100"],
            ["Channel", "PIM Marketplace", "Mapped to a BC Company Name, not a Shopify channel yet"],
            ["Export mapping", "Shopify Field on attribute", "title, vendor, body_html, metafields.* — preview map only"],
        ],
        [4.0, 5.5, 7.5],
    )

    heading(doc, "4.2 Marketplace", 2)
    body(
        doc,
        "Each marketplace (DE, AT, ES, CH, NP, CZ) stores the exact Business Central company name, Enabled, optional Template Item No. in the target company (posting groups on insert only), and flags to copy unit price or posting groups from master. "
        "PIM Item Marketplace is the per-item tick list: which countries receive this product.",
    )

    heading(doc, "5. Enrichment process", 1)
    body(doc, "1. Tell Me → PIM Families → Create default PIM setup (or rely on app install seed).")
    body(doc, "2. Tell Me → PIM Marketplaces → set Company Name (copy from Companies), Enabled, Template Item No. in the child, then Sync PIM setup to this company.")
    body(doc, "3. Open an Item. Set family DEFAULT (or another family), category, and Published to Webshop.")
    body(doc, "4. PIM Enrichment → Load family attributes → fill values. Completeness % uses required family attributes.")
    body(doc, "5. Tick Sync to companies, then run Sync to companies. Check PIM Sync Log.")
    body(doc, "Validating the family on the Item creates empty PIM Product Value rows so the enrichment page is ready immediately.")

    heading(doc, "6. Company sync", 1)
    add_picture(doc, imgs["sync"])
    caption(doc, "Figure 4 — Master to child sync. Same Item No. Template item is never used as the number.")

    heading(doc, "6.1 What is copied", 2)
    add_table(
        doc,
        ["Copied (product master)", "Not copied (ERP operations)"],
        [
            ["Description, search description, GTIN, tariff, origin, weights", "Inventory / stock"],
            ["Sales and purchase units of measure, item UOMs, master UOM", "Unit cost"],
            ["Item category, PIM family / category / published, PIM values", "Sales and purchase documents"],
            ["Picture (MediaSet), variants, translations", "Warehouse SKUs"],
            ["Extended text headers and lines", "Full TransferFields of Item"],
            ["Item references, document attachments (media and URL/URI)", "PTE Marketplace Code / Master SKU"],
            ["Standard Item Attributes matched by Name (IDs stay local)", ""],
            ["Optional: unit price; optional: posting groups from master", ""],
        ],
        [8.5, 8.5],
    )
    body(
        doc,
        "On insert, posting groups can come from a template item that already exists in the target company. "
        "Standard Item Attribute IDs must not be copied; each company generates its own IDs and matching is by Name.",
    )

    heading(doc, "6.2 Tenancy limit", 2)
    body(
        doc,
        "ChangeCompany only works inside one Business Central environment. If Germany (or any country) is a separate tenant or SaaS organisation, this sync cannot reach it. That requires APIs (not in this package).",
    )

    heading(doc, "7. Webshop preview", 1)
    add_picture(doc, imgs["webshop"])
    caption(doc, "Figure 5 — Control add-in flow. Catalog is filtered to PIM Published and capped at 48 items.")
    body(
        doc,
        "Product JSON includes PIM attributes, standard item attributes, Shopify field map, variants, UOMs, translations, extended texts, documents, and pictures. "
        "Control add-in scripts and CSS paths are relative to app.json. If the Webshop folder is under src, use src/Webshop/scripts/... and src/Webshop/styles/productViewer.css.",
    )

    heading(doc, "8. Object catalogue", 1)
    add_table(
        doc,
        ["ID", "Object", "Role"],
        [
            ["50600", "Enum PIM Attribute Type", "Text, Number, Yes/No, Option, Date"],
            ["50600–50606", "PIM tables", "Groups, attributes, options, family, family attrs, category, values"],
            ["50607", "Tableext PIM Item Ext", "Family, category, published on Item"],
            ["50632", "Codeunit PIM Enrichment", "Ensure family rows, completeness %"],
            ["50633", "Codeunit Product Visual Data", "Catalog and product JSON"],
            ["50634 / 50635", "Install / Upgrade", "Seed defaults per company"],
            ["50636–50638", "Marketplace, Item Marketplace, Sync Log", "Distribution model"],
            ["50639", "Codeunit PIM Company Sync", "ChangeCompany copy"],
            ["50610–50619", "PIM setup and enrichment pages", "Tell Me administration and lists"],
            ["50640–50642", "Marketplaces, item ticks, sync log pages", "Sync administration"],
            ["50351 / 50352", "Product Visual Card / Product Webshop", "Storefront pages"],
            ["50353 / 50354", "Item Card / Item List extensions", "PIM FastTab and actions"],
            ["50355", "Permission set PIM and Webshop", "Assignable"],
            ["—", "Control add-in ProductVisualViewer", "HTML/JS/CSS shop"],
        ],
        [3.4, 5.6, 8.0],
    )

    heading(doc, "9. Deployment and test", 1)
    body(doc, "Publish ICS Master (including the Webshop folder) in the master company and in every child company. Assign permission set PIM and Webshop if Tell Me is empty.")
    add_table(
        doc,
        ["Step", "Where", "Expected result"],
        [
            ["Download symbols, zero Problems, F5", "VS Code / AL", "App publishes"],
            ["Create default PIM setup", "PIM Families", "DEFAULT family, attributes, DE–CZ marketplaces"],
            ["Set Company Name + Enabled + template", "PIM Marketplaces", "Exact names from the Companies page"],
            ["Enrich one item, set Published", "Item Card / PIM Enrichment", "Completeness % moves"],
            ["Open Product Webshop", "Tell Me or Item actions", "Shopify-style catalog of published items"],
            ["Sync to one child, open same Item No.", "Child company", "Commercial data present; stock not copied"],
        ],
        [5.0, 4.5, 7.5],
    )

    heading(doc, "10. Design rules (do not break)", 1)
    rules = [
        "Item No. is the identity across countries.",
        "PIM is additive. Do not TransferFields the whole Item.",
        "Do not copy inventory, unit cost, or documents.",
        "Do not clash with PTE Marketplace Code or Master SKU.",
        "Match standard Item Attributes by Name, never by ID.",
        "Unpublished items must not appear in the catalog.",
        "Template Item No. supplies posting groups in the target only; it is never the product number.",
        "Webshop is an internal preview, not a customer storefront and not Shopify OAuth.",
    ]
    for i, rule in enumerate(rules, 1):
        body(doc, f"{i}. {rule}", space_after=4)

    heading(doc, "11. Later extensions (not in this package)", 1)
    body(doc, "REST/OData if countries are separate tenants. Locale-specific PIM values per marketplace. Push to Shopify Admin API using Shopify Field. Approval before publish. Catalog paging beyond 48 items.")

    heading(doc, "12. Glossary", 1)
    add_table(
        doc,
        ["Term", "Meaning in this solution"],
        [
            ["Master product company", "The BC company where PIM enrichment is done."],
            ["Child / marketplace", "A BC company (DE, AT, ES, CH, NP, CZ) that receives a copy of the item."],
            ["Family", "The set of attributes that apply to a product type."],
            ["Completeness", "Percentage of required family attributes that have a value."],
            ["Published", "Item.PIM Published — required for catalog visibility."],
            ["ChangeCompany", "AL API to read/write another company in the same environment."],
            ["Control add-in", "Embedded HTML/JS UI hosted on a BC page."],
        ],
        [4.5, 12.5],
    )

    body(doc, "This document is the architecture baseline for ICS Master PIM. Change tables, sync, or the control add-in against these rules.", bold=True)

    doc.save(str(DOC_PATH))
    print(f"Wrote {DOC_PATH} ({DOC_PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
